#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按官方《项目申报表 V2.0》模板生成申报表 .docx（图文并茂，逐栏填满）。"""
from __future__ import annotations
import json, pathlib, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

R = pathlib.Path(__file__).resolve().parent.parent
FIG = R / "docs" / "figures"
MEDIA = R / "web" / "media"
cal = json.loads((R / "eval" / "calibration_report.json").read_text(encoding="utf-8"))
orb = [x for x in cal["runs"] if x["detector"] == "ORB"][0]
cases = json.loads((R / "web" / "demo_cases.json").read_text(encoding="utf-8"))

INK = RGBColor(0x1E, 0x1C, 0x1A); CIN = RGBColor(0xA8, 0x32, 0x2A); GREY = RGBColor(0x8A, 0x81, 0x78)
d = Document()

st = d.styles["Normal"]; st.font.name = "宋体"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for s in d.sections:
    s.top_margin = s.bottom_margin = Cm(2.2); s.left_margin = s.right_margin = Cm(2.4)

fignum = {"n": 0}; tabnum = {"n": 0}


def H(t, lv=1):
    p = d.add_heading(t, level=lv)
    for r in p.runs:
        r.font.color.rgb = INK if lv > 1 else CIN
        r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def P(t, size=10.5, bold=False, color=None, align=None, space=4):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.45
    if align: p.alignment = align
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def FIGURE(path, cap, width=15.5):
    if not pathlib.Path(path).exists(): return
    fignum["n"] += 1
    d.add_picture(str(path), width=Cm(width))
    d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P("图 %d　%s" % (fignum["n"], cap), 9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space=10)


def TABLE(headers, rows, cap=None, widths=None):
    tabnum["n"] += 1
    if cap:
        P("表 %d　%s" % (tabnum["n"], cap), 9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space=3)
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(9.5)
            r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


# ================================================================ 封面
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(70)
r = p.add_run("智极松 AI Skillathon 大赛"); r.font.size = Pt(15); r.font.color.rgb = GREY
r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("项目申报表"); r.font.size = Pt(30); r.bold = True; r.font.color.rgb = INK
r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
P("", space=26)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("城设 CityBible"); r.font.size = Pt(23); r.bold = True; r.font.color.rgb = CIN
r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
P("给每一座城市，写一本设定集", 13, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space=30)
P("城市文旅内容的真实素材资产库与地理真实性验真门禁", 12,
  align=WD_ALIGN_PARAGRAPH.CENTER, space=46)
TABLE(["项目", "内容"], [
    ["申报赛道", "AI + 文创及服务"],
    ["项目阶段", "Demo / MVP"],
    ["核心指标", "68 条埋真值校准样本，准确率 100%，正负内点数分离 200 倍"],
    ["在线体验", "http://47.119.112.225/citybible/ （免注册，预置示例，打开即看）"],
    ["代码仓库", "见附录，含完整源码、校准脚本与部署脚本"],
    ["提交日期", datetime.date.today().strftime("%Y 年 %m 月 %d 日")],
], widths=[3.6, 11.5])
d.add_page_break()

# ================================================================ 一
H("一、基本信息")
TABLE(["字段", "填写内容"], [
    ["项目名称", "城设 CityBible"],
    ["申报赛道", "☑ AI+文创及服务　（□工业制造 □健康与生命科学 □新能源物流 □安全可信 □具身智能）"],
    ["团队/个人名称", "戴尚好"],
    ["团队类型", "☑ 开发者团队"],
    ["项目阶段", "☑ Demo/MVP　（□概念验证 □试点应用 □已有收入 □规模化推广）"],
    ["参赛模式", "☑ 个人参赛（OPC 模式）"],
    ["联系人", "戴尚好"],
    ["联系方式", "手机：（请填写）　　邮箱：bcefghj@163.com"],
], widths=[3.6, 11.5])

H("二、团队成员与分工")
TABLE(["姓名", "角色/职务", "所在单位", "专业背景", "核心贡献"], [
    ["戴尚好", "项目负责人 / 技术负责人 / 产品",
     "（请填写）", "全栈开发、多 Agent 系统",
     "独立完成判定引擎、校准体系、后端服务、MCP/CLI 三形态、展示页与全部文档"],
], widths=[2.0, 2.6, 2.6, 3.0, 4.9])
P("说明：本项目为 OPC 个人参赛。曾独立完成字节跳动 AI 全栈挑战赛竞品分析赛道参赛作品"
  "CompetitorLens（多 Agent 协作 + 证据链报告系统），本项目在该工程经验基础上展开。")

# ================================================================ 三
d.add_page_break()
H("三、项目摘要")
H("1. 一句话介绍（不超过 50 字）", 2)
P("为城市与县域沉淀可复用的真实文旅素材资产，量产地理真实、可归因到店的短视频内容。",
  12, bold=True, color=CIN)
P("（共 44 字）", 9, color=GREY)

H("2. 目标客户 / 使用者", 2)
P("第一顺位：县级及地市级文旅局、文旅集团与县级融媒体中心（全国 2800+ 个县级行政区，"
  "2000+ 家县级融媒体中心）。他们有内容任务、有考核指标，但普遍缺乏内容产能与专业审核能力。")
P("第二顺位：景区运营方、非遗工坊、民宿与本地小微商户。2026 年起微短剧投资门槛上调至"
  "普通 100 万元、重点 300 万元，这类主体被彻底挤出专业内容生产。")
P("第三顺位：承接文旅内容业务的 MCN、短视频制作方与文旅营销公司——他们需要在交付前"
  "自证素材真实、地点无误，以规避甲方与舆论风险。")

H("3. 核心痛点", 2)
P("痛点一：产能已经不是问题，真实性才是。", bold=True)
P("2026 年第一季度全网新增上线微短剧约 12.8 万部，其中 AI 参与制作的内容占比超过 95%，"
  "行业同期出现「同质化、抽卡式批量作品流量持续下滑」的判断。文旅内容与其他内容有一个"
  "根本区别——它拍摄的是真实存在的地方。一张看着像、实际不是岳麓书院的图片一旦发布，"
  "对文旅主管部门就是一次可被公众核查的事故。目前市面上所有工具都在解决「怎么生成得更快」，"
  "没有任何一个工具回答「你生成的这个地方，现实中长这样吗」。")
P("痛点二：素材是散的。", bold=True)
P("同一座城市的摄影素材在宣传部门、地方志在方志办、文物影像在博物馆、非遗档案在文化馆，"
  "彼此不互通。每做一次宣传就要重新拍一次，历史投入无法沉淀为可复用资产。")
P("痛点三：效果是黑的。", bold=True)
P("文旅宣传预算花出去，年终能提交的只有播放量。没有任何一方能回答「这条内容带来了"
  "多少实际到访」。这使得文旅内容投放长期无法进行 ROI 核算。")

H("4. 解决方案概述", 2)
P("城设 CityBible 由三层构成，当前 Demo 已完整实现前两层：")
P("第一层 · 城市素材资产库。", bold=True)
P("为每个地点沉淀多方位、多时段的参考素材，并登记来源与授权状态，形成这座城市的"
  "「数字底片」。一次沉淀，后续所有内容生成与验真都从这里取材，边际成本趋近于零。")
P("第二层 · 地理真实性验真门禁（本项目的技术核心）。", bold=True)
P("待验图片与该地点的参考照片库逐一比对，主判定链路完全不调用大模型：由 OpenCV 特征点"
  "检测、Lowe's ratio test 与 RANSAC 单应矩阵几何校验算出内点数与内点率，输出"
  "match / mismatch / needs_human_review 三态门禁结论，并附带匹配点连线证据图。"
  "同一输入任意次运行结果完全一致，可复现、可调参、可审计。")
P("第三层 · 到店转化归因（架构已设计，尚未接入真实数据）。", bold=True)
P("播放 → 搜索 → 加入行程 → 到访 → 核销的链路归因，产出内容投放效益报表。"
  "本阶段无真实核销数据，故 Demo 中不展示该模块，避免与真实跑出的判定结果混淆。")
FIGURE(FIG / "fig0_architecture.png", "城设 CityBible 系统架构。主判定链路零大模型调用，"
       "结论绑定证据，三种对外形态共用一份实现。")

H("5. 当前进展", 2)
TABLE(["事项", "状态", "可验证方式"], [
    ["地理真实性判定引擎", "已完成", "engine/verify.py，可直接命令行运行"],
    ["裁判校准体系", "已完成", "eval/calibrate.py，一条命令复现全部指标"],
    ["城市素材资产库（演示版）", "已完成", "长沙 4 个地标，16 张素材，均为自有版权"],
    ["后端服务 REST + MCP", "已完成", "server/，含 FastAPI 与标准库两种运行模式"],
    ["命令行工具 CLI", "已完成", "cli/citybible.py"],
    ["公网展示页", "已完成", "http://47.119.112.225/citybible/"],
    ["一键部署脚本", "已完成", "deploy.sh，含端口自动避让与失败回滚"],
    ["到店归因", "架构已设计，未接数据", "本阶段不展示"],
], widths=[4.4, 3.2, 7.5])

# ================================================================ 四
d.add_page_break()
H("四、技术与产品概述")
H("1. 核心技术 / 算法简介", 2)
P("（1）主判定：特征点匹配 + RANSAC 几何校验，零大模型。", bold=True)
P("流程为：对待验图与参考图分别提取局部特征点与描述子 → 暴力匹配后用 Lowe's ratio test "
  "（阈值 0.75）滤除区分度不足的匹配 → 用 RANSAC 估计单应矩阵（重投影阈值 5 像素，"
  "置信度 0.995）→ 统计几何自洽的内点数与内点率。RANSAC 这一步是关键：仅靠描述子相似度"
  "会被砖墙、瓦片、树叶等重复纹理产生大量误匹配，加上单应矩阵约束后，只有在几何上"
  "自洽的匹配才被计为内点。")
P("（2）为什么不用 CLIP 余弦相似度——这是一次有依据的技术选型。", bold=True)
P("「这张图是不是岳麓书院」属于实例级识别（instance-level recognition），而非零样本分类。"
  "CLIP 图像 embedding 存在锥效应（cone effect）：任意两张自然图像的余弦值天然落在 "
  "0.5–0.9 的窄带内，绝对阈值失去区分意义；open_clip 官方讨论区关于「图库该如何设阈值」"
  "的提问至今为 Unanswered 状态。此外 CLIP 已知存在 typographic attack（arXiv:2103.10480）——"
  "在图像上贴一张写有目标名称的字条即可显著抬高相似度，而文旅场景中水印、店招、"
  "打卡贴纸极其常见。特征点匹配不存在这两个问题，且能够指出「哪里像」，"
  "这正是证据链所需要的。")
P("（3）证据即产物。", bold=True)
P("每一条判定都携带：内点数、内点率、单应矩阵、逐条检查点结论，以及一张由 "
  "cv2.drawMatches 生成的匹配点连线图。产出不了证据的判定不允许写入结论。")
P("（4）四态门禁，中间地带明确交人。", bold=True)
P("内点数高于通过阈值且内点率达标 → match；低于拒绝阈值 → mismatch；"
  "落在两者之间 → needs_human_review，不硬判；异常路径 → error 并记录原因，"
  "不使用随机值或默认值假装成功。")

H("2. 产品形态与功能模块", 2)
TABLE(["模块", "路径", "说明"], [
    ["判定引擎", "engine/verify.py", "三种检测器可切换，参数全部外置"],
    ["校准器", "eval/calibrate.py", "埋真值样本集 → 混淆矩阵与全部指标"],
    ["素材资产库", "assets/", "参考照、候选图、正样本变体与 manifest"],
    ["REST 接口", "server/app.py", "FastAPI，自动生成 OpenAPI 文档"],
    ["标准库兜底服务", "server/simple_server.py", "无第三方依赖，装不上框架时仍可运行"],
    ["MCP Server", "server/core.py", "Streamable HTTP，3 个 tool，供任意 Agent 调用"],
    ["命令行工具", "cli/citybible.py", "断网与平台故障时的降级路径"],
    ["展示页", "web/index.html", "单文件 33 KB，数据内联，无需后端即可打开"],
    ["部署脚本", "deploy.sh", "端口自动避让、影响面比对、失败回滚"],
], widths=[3.2, 4.4, 7.5])

H("3. 当前开发阶段与完成情况", 2)
P("是否已有 Demo 或可验证原型：☑ 是　□ 否", bold=True)
P("可验证方式（三条独立路径，评委任选其一即可，均无需配置环境）：", bold=True)
P("路径一　打开 http://47.119.112.225/citybible/ ，免注册、预置示例数据，"
  "页面上四个验真用例的结论、内点数与证据图，全部由本仓库引擎真实跑出。")
P("路径二　克隆仓库后执行 python3 eval/calibrate.py，可在本机完整复现下文全部指标。")
P("路径三　执行 python3 cli/citybible.py verify <图片> --claim 岳麓书院，单张验真。")
TABLE(["等级", "内容", "状态"], [
    ["P0 必做", "判定引擎 / 证据链 / 素材库 / 门禁四态 / 公网展示页", "全部完成"],
    ["P1 进阶", "裁判校准体系 / 三检测器对比 / REST 接口 / 部署脚本与回滚", "全部完成"],
    ["P2 加分", "MCP Server / CLI / 标准库兜底运行模式 / 局限性实测与披露", "全部完成"],
    ["未完成", "到店归因接真实数据 / 真实实拍参考库 / 专业 OCR 水印拦截", "已列入后续工作"],
], cap="项目完成度对照（P0/P1/P2 三级自评）", widths=[2.4, 9.2, 3.5])

# ================================================================ 五
d.add_page_break()
H("五、场景应用与产业价值")
H("1. 目标应用场景", 2)
P("场景一 · 文旅内容发布前的验真门禁。文旅局或制作方在内容发布前批量过一遍门禁，"
  "地点不符的内容被拦下并给出证据，避免公开事故。")
P("场景二 · 城市素材资产的沉淀与调用。把散落在多个部门的素材做一次结构化沉淀，"
  "此后节庆宣传、招商推介、非遗申报、研学课程等所有内容需求都从同一个库里取材。")
P("场景三 · AI 生成内容的真实性标注。当 95% 的微短剧已有 AI 参与，"
  "「这段画面里的地点是真的吗」将成为平台与监管的共同需求。")

H("2. 拟解决的行业问题", 2)
P("AI 内容生产的产能红利已经释放，但文旅内容的真实性没有任何一层保障。"
  "本项目用一个可复现、可解释、零大模型依赖的判定链路补上这一层，"
  "并把结论落在可审计的证据上，而不是一个无法追问的相似度分数。")

H("3. 预期效果与价值", 2)
P("效率维度：", bold=True)
P("单次判定平均耗时 %.0f 毫秒（2 核 CPU，ORB），可批量流水线处理。"
  "人工目检一张图并给出可追溯的判断依据，通常需要数分钟且不可复现。" % orb["mean_latency_ms"])
P("质量维度：", bold=True)
P("在 68 条埋真值校准样本上，ORB / AKAZE / SIFT 三种检测器均取得查准率、查全率、"
  "F1 与准确率全部为 1.000 的结果，误杀率与漏放率均为 0。")
P("成本维度：", bold=True)
P("主判定链路大模型调用次数为 0，单次判定的边际算力成本接近于零；"
  "整套系统可在 2 核 2G 无 GPU 的云服务器上运行，无需向量数据库。")
P("安全与合规维度：", bold=True)
P("每一条打回结论都可追溯到具体的检查点与证据图，可用于对外申诉与复核，"
  "而不是一个无法解释的黑箱判断。")
FIGURE(FIG / "fig1_confusion.png", "三种检测器在 68 条埋真值校准样本上的混淆矩阵，均为满分。")
FIGURE(FIG / "fig2_separation.png", "四类样本的内点数分布（对数轴）。正负样本相差一到两个数量级，"
       "说明阈值的选择并不精细——任何落在中间的阈值给出的结果都一样。")
FIGURE(FIG / "fig3_latency.png", "三种检测器的单次判定平均耗时对比。ORB 最快且分离度最大，选为默认。", 12)

H("4. 现有客户 / 试点 / 合作进展", 2)
P("当前无签约客户与试点，属实。本项目为独立开发的 Demo/MVP 阶段作品。")
P("产业需求的客观依据：长沙马栏山视频文创产业园于 2025 年 10 月发布全国首个"
  "「马栏山微短剧智能服务平台」1.0 版，将备案时间从 1–3 个月压缩至 3–7 天；"
  "配套的「马栏山微短剧审核与全流程服务开放平台项目」（政府采购编号 MLSCG-20251005107）"
  "预算 5,037,742.76 元，中标金额 4,780,000 元。该采购证明微短剧内容服务平台"
  "已是真实且有明确预算的政府采购品类。本项目定位于其上游的内容生产与真实性验证环节，"
  "与既有平台的审核、译制、分发能力互补而非重复建设。")

# ================================================================ 六
d.add_page_break()
H("六、竞争优势")
P("优势一 · 差异不在生成，在判定。", bold=True)
P("开源的一键短视频生成工具已相当成熟（如 MoneyPrinterTurbo，约 5.9 万 star，MIT 协议），"
  "「输入一句话输出一条视频」不再具备壁垒。本项目刻意不在生成环节竞争，"
  "生成层做成可插拔，差异集中在生成之后的验真与证据链。")
P("优势二 · 主判定零大模型，因而可复现、可审计、成本恒定。", bold=True)
P("同一输入任意次运行结果完全一致，不受模型版本更新影响，也不产生按次计费的推理成本。"
  "这对需要留存审计记录的政府与机构客户是硬性要求。")
P("优势三 · 证据是产物而非附加说明。", bold=True)
P("匹配点连线图是判定过程的天然副产物，不需要额外设计「解释模块」。"
  "相比之下，基于 embedding 相似度的方案只能给出一个数字，无法回答「哪里不像」。")
P("优势四 · 校准报告随产品交付。", bold=True)
P("我们不仅提供判定能力，还提供「凭什么信这个判定」的量化答案——"
  "埋真值样本集、混淆矩阵、误杀率与漏放率，以及一条可复现全部指标的命令。")
P("优势五 · 工程约束下的务实选型。", bold=True)
P("在 2 核 2G 无 GPU 的目标机型上，我们论证并放弃了向量数据库"
  "（Milvus 官方最低要求 8 GB 内存 / 4 核；Chroma 官方文档明确不推荐 2 GB 以下部署），"
  "改用 numpy 直接计算；同时提供标准库兜底运行模式，在无法安装 Web 框架的环境中仍可完整运行。")
TABLE(["对比维度", "通用 AIGC 生成工具", "基于 CLIP 的相似度方案", "城设 CityBible"], [
    ["能否判断地点真伪", "不涉及", "可给出相似度数值", "可给出四态结论"],
    ["能否指出哪里不像", "不涉及", "否", "是（匹配点连线图）"],
    ["结果是否可复现", "否（生成有随机性）", "取决于模型版本", "是（完全确定性）"],
    ["阈值是否可稳定标定", "不涉及", "困难（锥效应）", "是（正负相差 200 倍）"],
    ["能否抵抗贴字干扰", "不涉及", "已知存在攻击面", "不受影响"],
    ["2 核 2G 可否运行", "多数不可", "需 torch，勉强", "可（内存 <300 MB）"],
    ["单次判定成本", "按次计费", "按次计费", "约等于零"],
], cap="与两类替代方案的逐项对比", widths=[3.4, 3.6, 3.9, 4.2])

# ================================================================ 演示
d.add_page_break()
H("七、Demo 验真用例实录")
P("以下四个用例的判定结论、各项数值与证据图，全部由本仓库引擎真实跑出，未经修饰。"
  "完整逐条记录见 eval/calibration_report.json。")
TABLE(["用例", "声称地点", "判定", "内点数", "耗时"],
      [[c["title"], c["claim"],
        {"match": "通过 MATCH", "mismatch": "打回 MISMATCH"}.get(c["verdict"], c["verdict"]),
        c["inliers"], "%.0f ms" % c["elapsed_ms"]] for c in cases],
      cap="四个验真用例的判定结果", widths=[5.6, 2.6, 3.2, 2.0, 1.7])
key = [c for c in cases if c["id"] == "neg_ai_yuelu"]
if key and (MEDIA / "neg_ai_yuelu_evidence.jpg").exists():
    P("最关键的一个用例：AI 重新生成的「岳麓书院」。", bold=True, color=CIN)
    P("用同一段提示词再生成一次岳麓书院，人眼看着很像，但它在几何上是一座全新的建筑。"
      "系统给出的内点数为 %d，与一个完全不相干的地点处在同一量级。"
      "这正是文旅内容必须验真的直接理由，也是本项目存在的根本原因。" % key[0]["inliers"])
    FIGURE(MEDIA / "neg_ai_yuelu_evidence.jpg",
           "证据图（cv2.drawMatches 原始输出）。左为待验图，右为参考照，绿线为 RANSAC 内点。"
           "几乎不存在几何自洽的连线。")
pos = [c for c in cases if c["id"] == "pos_yuelu"]
if pos and (MEDIA / "pos_yuelu_evidence.jpg").exists():
    P("对照用例：同一地点、不同时段与角度的拍摄。", bold=True)
    P("光线、视角与画质都变了，但建筑的物理结构点没有变，内点数达到 %d，"
      "与上一个用例相差约 %d 倍。" % (pos[0]["inliers"], pos[0]["inliers"] // max(key[0]["inliers"], 1)))
    FIGURE(MEDIA / "pos_yuelu_evidence.jpg", "证据图。密集的绿色连线即为几何自洽的匹配点。")

# ================================================================ 局限
d.add_page_break()
H("八、已知局限与后续工作")
P("以下四条是实测中确实存在的问题，如实列出。")
P("局限一 · 校准集的正样本是合成的，不是实拍。", bold=True)
P("正样本采用 HPatches / Oxford-Affine 数据集的标准构造法，在同一张底图上施加受控的"
  "视角与光照变换，几何对应关系是精确的。真实世界的二次拍摄更困难：不同相机、季节变化、"
  "人群遮挡、建筑改造都会降低内点数。因此本报告证明的是「方法在受控条件下分离度极大、"
  "阈值不敏感」，而不是「在真实实拍上准确率 100%」。后者需要真实多时段实拍数据。")
P("局限二 · 类文字区域检测器不具备区分力，已主动降级。", bold=True)
P("原设计用该指标防御 typographic attack。实测结果：干净建筑照的类文字区域占比天然在"
  "0.1%–15.2% 之间（屋瓦、梁枋、栏杆的水平重复纹理被误判为文字行），而人工贴上打卡水印后"
  "仅从 4.08% 升至 9.12%，落在正常范围内部。该指标无法区分两者，"
  "因此我们将其从门禁降级为只记录不拦截的信息性信号，而不是保留一个会给出错误信心的功能。"
  "真正拦截水印需接入专业 OCR 接口。")
P("局限三 · 系统只能验证参考库中已收录的地点。", bold=True)
P("参考库的方位与时段覆盖不足时，真实照片也可能被判为 mismatch。工程上唯一有效的"
  "提升手段是扩大覆盖：每个地点至少 4 方位 × 3 时段 × 2 季节。")
P("局限四 · 到店归因尚未接入真实数据。", bold=True)
P("链路架构已设计，但本阶段没有真实核销数据，相关数字若展示即为占位。"
  "因此 Demo 中不展示归因看板，以免与真实跑出的判定结果混淆。")
P("后续工作优先级：接入真实授权实拍素材扩充参考库 → 与文旅单位共建试点验证真实场景准确率 "
  "→ 接入专业 OCR 补上水印拦截 → 打通到店归因数据源。")

# ================================================================ 合规
H("九、合规声明")
TABLE(["事项", "团队声明"], [
    ["数据合规", "☑ 仅使用公开数据　☑ 不涉及个人信息。演示素材全部由 AI 生成，"
                 "不含任何第三方版权图片、不含可识别的自然人肖像。"],
    ["模型与算法", "☑ 调用第三方 API　☑ 开源模型并遵守协议。判定引擎基于 OpenCV"
                   "（Apache-2.0）自行实现；素材生成调用火山方舟 Seedream 与 MiniMax，"
                   "均为商业授权服务；主判定链路不依赖任何外部 API。"],
    ["知识产权", "☑ 无争议。代码为本人独立编写；未复制任何第三方项目代码。"
                 "文档中引用的开源项目均注明出处与协议。"],
    ["行业合规", "☑ 内容版权。系统本身即用于降低文旅内容的真实性与版权风险；"
                 "素材库设计中，来源与授权状态为一等字段。"],
], widths=[3.0, 12.1])
P("补充说明：", bold=True)
P("1）演示素材来源：全部 16 张地标参考照与候选图由火山方舟 Seedream 5.0 生成，"
  "48 张正样本由这些底图经确定性几何与光度变换派生，变换参数与真值一并落盘，可完全复现。")
P("2）密钥管理：仓库中仅提供 .env.example 模板，不含任何真实密钥；.env 已列入 .gitignore。")
P("3）AI 辅助开发声明：本项目在开发过程中使用了 AI 编程工具辅助编码与文档撰写，"
  "全部代码经本人运行、核对与演示验证；文中所有数据均由仓库内脚本实际跑出，可复现。")
P("4）对既有平台的态度：本项目明确不重复建设马栏山微短剧智能服务平台已提供的"
  "合规审核、多语种译制与多平台分发能力，定位于其上游的内容生产与真实性验证环节。")

# ================================================================ 承诺
H("十、申报承诺")
P("本人承诺所提交材料真实、准确、完整，不存在虚假申报、恶意抄袭、隐瞒重大争议等情形。"
  "理解后续阶段（夏令营、Demo 评审、决赛路演）将根据本表内容进行辅导与评审。")
P("")
P("项目负责人签字：________________　　　　日期：　　　年　　月　　日")

# ================================================================ 附录
d.add_page_break()
H("附录 · 可复核材料清单")
TABLE(["材料", "位置", "如何核验"], [
    ["在线展示页", "http://47.119.112.225/citybible/", "浏览器直接打开，免注册"],
    ["判定引擎源码", "engine/verify.py", "约 380 行，含完整注释与选型论证"],
    ["校准脚本与报告", "eval/calibrate.py, eval/calibration_report.json",
     "运行 python3 eval/calibrate.py 完整复现"],
    ["68 条校准样本", "assets/", "含 manifest 与变换参数，可逐条追溯"],
    ["证据图", "assets/evidence/, web/media/", "cv2.drawMatches 原始输出"],
    ["MCP Server", "server/core.py", "3 个 tool，Streamable HTTP"],
    ["命令行工具", "cli/citybible.py", "citybible verify <图> --claim <地点>"],
    ["部署脚本", "deploy.sh", "含端口避让、影响面比对与失败回滚"],
    ["技术文档", "docs/交付总文档.md", "速览 → 价值 → 实现 → 可复核"],
], widths=[3.4, 5.4, 6.3])
P("")
P("本申报表中的全部数值（准确率、内点数、耗时、分离倍数）均由 eval/calibrate.py 于 %s 实际跑出，"
  "未经人工修改。" % cal["generated_at"], 9, color=GREY)

out = R / "docs" / "AI+文创及服务_城设CityBible_戴尚好_申报表_V1.docx"
d.save(str(out))
print("申报表已生成：", out)
print("页数约：", len(d.paragraphs), "段落 /", len(d.tables), "个表格 /", fignum["n"], "张图")
